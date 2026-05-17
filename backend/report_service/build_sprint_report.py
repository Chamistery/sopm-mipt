#!/usr/bin/env python3
"""
Генератор отчётов об итерации для СУПП ВШПИ МФТИ в двух форматах:

    build_docx(data)     — DOCX для подписи (командный) / build_student_docx — личный
    build_markdown(data) — полный сырой MD со всей информацией по спринту

Запуск без аргументов создаёт 8 файлов:
    sprint_report_{template,example}{,_student}.{docx,md}

DOCX-синтаксис плейсхолдеров: Mustache-like {{field}} и {{#list}}…{{/list}}.

Правила оформления DOCX:
- шрифт Times New Roman 12pt везде, кроме нумерованных заголовков 1-го уровня (14pt bold);
- выравнивание «по ширине» везде, кроме центрированных блоков (МФТИ-шапка, титул,
  подзаголовок, заголовки таблиц);
- ширины столбцов подобраны без пустого места: полезная ширина страницы 17 см.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Полезная ширина страницы A4 при полях 2.5 / 1.5 см = 17 см
PAGE_W = 17.0


# ════════════════════════════════════════════════════════════
#                         DOCX helpers
# ════════════════════════════════════════════════════════════

def set_cell_white(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFFFF")
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))  # 567 twips ≈ 1 cm
    tcW.set(qn("w:type"), "dxa")
    tc_pr.append(tcW)


def _apply_para_format(p):
    """Единые правила абзаца: интервалы до/после 0, межстрочный 1.5."""
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5


def _align_paragraphs(cell, alignment):
    for p in cell.paragraphs:
        p.alignment = alignment
        _apply_para_format(p)


def _size_runs(cell, size=12, bold=None):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size)
            if bold is not None:
                r.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _apply_para_format(p)
    run = p.add_run(text)
    run.bold = True
    # Level 1 — нумерованные разделы: 14pt.
    # Level 2 — подзаголовки (1.1, 1.2, ФИО участника): 12pt.
    run.font.size = Pt(14 if level == 1 else 12)
    return p


def add_field(doc, label, value):
    """Строка «Метка: значение» — 12pt, по ширине."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _apply_para_format(p)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = p.add_run(value if value is not None else "—")
    r2.font.size = Pt(12)


def add_text(doc, text):
    p = doc.add_paragraph(text or "—")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _apply_para_format(p)
    for r in p.runs:
        r.font.size = Pt(12)


def make_table(doc, headers, widths):
    """Шапка таблицы: bold + центрирование, 12pt, белый фон."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_white(hdr[i])
        set_cell_width(hdr[i], widths[i])
        _align_paragraphs(hdr[i], WD_ALIGN_PARAGRAPH.CENTER)
        _size_runs(hdr[i], size=12, bold=True)
    return table


def add_row(table, values, widths, justify_text=True):
    """Строка данных: 12pt, текстовые ячейки — по ширине (justify)."""
    row = table.add_row().cells
    for i, v in enumerate(values):
        row[i].text = str(v) if v is not None else "—"
        set_cell_white(row[i])
        set_cell_width(row[i], widths[i])
        _align_paragraphs(row[i], WD_ALIGN_PARAGRAPH.JUSTIFY if justify_text
                          else WD_ALIGN_PARAGRAPH.LEFT)
        _size_runs(row[i], size=12)
    return row


def setup_document(doc):
    """Базовый стиль: TNR 12pt, justify, интервалы 0, межстрочный 1.5."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(1.5)


def add_mfti_header(doc):
    for line in [
        "Министерство науки и высшего образования Российской Федерации",
        "Федеральное государственное автономное образовательное учреждение",
        "высшего образования",
        "«Московский физико-технический институт",
        "(национальный исследовательский университет)»",
        "Высшая школа программной инженерии",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_para_format(p)
        run = p.add_run(line)
        run.font.size = Pt(12)
    p = doc.add_paragraph()
    _apply_para_format(p)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_para_format(p)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_para_format(p)
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(12)


# ── Ширины колонок (сумма = 17 см полезной ширины) ──
#  Состав команды: № | ФИО | Роль в команде
WIDTHS_MEMBERS = [1.0, 7.0, 9.0]

#  Задачи: № и Часы/Сроки — строго под заголовки, Задача = Описание
WIDTHS_TASKS = [0.9, 6.35, 1.6, 1.8, 6.35]

#  Встречи: Дата | Тема | Итоги — без пустого места
WIDTHS_MEETINGS = [2.5, 5.5, 9.0]

#  Сводные оценки: ФИО | Оценка / 10 | Комментарий
WIDTHS_SCORES = [6.5, 2.0, 8.5]


# ════════════════════════════════════════════════════════════
#                   DOCX — командный отчёт
# ════════════════════════════════════════════════════════════

def _render_team_body(doc, data):
    """Дозаписывает один командный отчёт в существующий Document.

    Не вызывает setup_document — вызывающая сторона делает это один раз
    на документ (см. build_docx и build_multi_docx).
    """
    add_mfti_header(doc)
    add_title(doc,
              f"ОТЧЁТ ОБ ИТЕРАЦИИ № {data['sprint']['number']}",
              "по дисциплине «Проектный практикум»")

    # 1. Сведения о проекте
    add_heading(doc, "1. Сведения о проекте", level=1)
    add_field(doc, "Проект", data["project"]["title"])
    add_field(doc, "Инициатор", data["project"]["company"])
    add_field(doc, "Ментор", data["project"]["mentor"])
    add_field(doc, "Команда", data["team"]["name"])
    add_field(doc, "Тимлид", data["team"]["leader"])
    add_field(doc, "Период итерации",
              f"с {data['sprint']['start_date']} по {data['sprint']['end_date']}")
    add_field(doc, "Дата формирования отчёта", data["report"]["generated_at"])

    add_heading(doc, "Состав команды", level=2)
    t = make_table(doc, ["№", "ФИО", "Роль в команде"], WIDTHS_MEMBERS)
    for i, m in enumerate(data["team"]["members"], 1):
        add_row(t, [i, m["full_name"], m["role_in_team"]], WIDTHS_MEMBERS)

    # 2. Командный отчёт
    add_heading(doc, "2. Командный отчёт", level=1)
    add_heading(doc, "2.1. Что сделано", level=2)
    add_text(doc, data["team_report"]["summary"])
    add_heading(doc, "2.2. Проблемы и риски", level=2)
    add_text(doc, data["team_report"]["problems"])
    add_heading(doc, "2.3. План на следующую итерацию", level=2)
    add_text(doc, data["team_report"]["next_plan"])

    # 3. Индивидуальный вклад — без колонок «Статус» и «MR»
    add_heading(doc, "3. Индивидуальный вклад участников", level=1)
    task_headers = ["№", "Задача", "Часы", "Сроки", "Описание работы"]
    for member in data["members_with_tasks"]:
        add_heading(doc, f"{member['full_name']} — {member['role_in_team']}", level=2)
        t = make_table(doc, task_headers, WIDTHS_TASKS)
        for i, task in enumerate(member["tasks"], 1):
            add_row(t, [
                i, task["name"], task["hours_estimate"],
                f"{task['start_date']} – {task['end_date']}",
                task["work_description"],
            ], WIDTHS_TASKS)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _apply_para_format(p)
        r = p.add_run("Оценка ментора за итерацию: ")
        r.bold = True
        r.font.size = Pt(12)
        r = p.add_run(f"{member['sprint_score']['score']} / 10")
        r.bold = True
        r.font.size = Pt(12)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run("Комментарий ментора: ")
        r.bold = True
        r.font.size = Pt(12)
        r = p.add_run(member["sprint_score"]["comment"])
        r.font.size = Pt(12)

    # 4. Встречи — без колонок «Время» и «Длительность»
    add_heading(doc, "4. Встречи с ментором в ходе итерации", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Согласно п. 8.5 Положения — не реже одного раза за итерацию.")
    r.italic = True
    r.font.size = Pt(12)

    t = make_table(doc, ["Дата", "Тема", "Итоги"], WIDTHS_MEETINGS)
    for m in data["meetings"]:
        add_row(t, [m["meeting_date"], m["title"], m["summary"]], WIDTHS_MEETINGS)

    # 5. Заключение ментора
    add_heading(doc, "5. Заключение ментора по итерации", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Комментарий к командному отчёту: ")
    r.bold = True
    r.font.size = Pt(12)
    r = p.add_run(data["team_report"]["mentor_comment"])
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _apply_para_format(p)
    r = p.add_run("Сводная таблица индивидуальных оценок:")
    r.bold = True
    r.font.size = Pt(12)

    t = make_table(doc, ["ФИО участника", "Оценка / 10", "Комментарий"], WIDTHS_SCORES)
    for m in data["members_with_tasks"]:
        add_row(t, [m["full_name"], m["sprint_score"]["score"],
                    m["sprint_score"]["comment"]], WIDTHS_SCORES)

    # Подписи
    sp = doc.add_paragraph()
    _apply_para_format(sp)
    p = doc.add_paragraph()
    _apply_para_format(p)
    r = p.add_run("Подписи:")
    r.bold = True
    r.font.size = Pt(12)

    sig_w = [5.5, 5.75, 5.75]
    tbl = doc.add_table(rows=0, cols=3)
    tbl.autofit = False
    for label, name, sign in [
        ("Тимлид команды", data["team"]["leader"], "____________________"),
        ("Ментор проекта", data["project"]["mentor"], "____________________"),
        ("Дата", data["report"]["generated_at"], ""),
    ]:
        cells = tbl.add_row().cells
        cells[0].text = label
        cells[1].text = name
        cells[2].text = sign
        for j, c in enumerate(cells):
            set_cell_width(c, sig_w[j])
            for para in c.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_para_format(para)
                for r in para.runs:
                    r.font.size = Pt(12)
                    if j == 0:
                        r.bold = True


def build_docx(data):
    """Один командный отчёт за один спринт."""
    doc = Document()
    setup_document(doc)
    _render_team_body(doc, data)
    return doc


def build_multi_docx(data_list):
    """Несколько командных отчётов в одном документе (page break между)."""
    if not data_list:
        raise ValueError("data_list must contain at least one sprint")
    doc = Document()
    setup_document(doc)
    for i, data in enumerate(data_list):
        if i > 0:
            doc.add_page_break()
        _render_team_body(doc, data)
    return doc


# ════════════════════════════════════════════════════════════
#             DOCX — индивидуальный отчёт студента
# ════════════════════════════════════════════════════════════

def build_student_docx(data, student_index=0):
    member = data["members_with_tasks"][student_index]
    team_member = next(
        (m for m in data["team"]["members"] if m["full_name"] == member["full_name"]),
        {},
    )

    doc = Document()
    setup_document(doc)
    add_mfti_header(doc)
    add_title(doc,
              f"ИНДИВИДУАЛЬНЫЙ ОТЧЁТ ОБ ИТЕРАЦИИ № {data['sprint']['number']}",
              "по дисциплине «Проектный практикум»")

    # 1. Сведения
    add_heading(doc, "1. Сведения", level=1)
    add_field(doc, "Участник", member["full_name"])
    add_field(doc, "Функциональная роль", member["role_in_team"])
    if team_member.get("group"):
        add_field(doc, "Группа",
                  f"{team_member['group']}, курс {team_member.get('course', '—')}")
    add_field(doc, "Проект", data["project"]["title"])
    add_field(doc, "Инициатор", data["project"]["company"])
    add_field(doc, "Ментор", data["project"]["mentor"])
    add_field(doc, "Команда", data["team"]["name"])
    add_field(doc, "Тимлид", data["team"]["leader"])
    add_field(doc, "Период итерации",
              f"с {data['sprint']['start_date']} по {data['sprint']['end_date']}")
    add_field(doc, "Дата формирования отчёта", data["report"]["generated_at"])

    # 2. Задачи
    add_heading(doc, "2. Выполненные задачи", level=1)
    if not member["tasks"]:
        add_text(doc, "Задач нет.")
    else:
        task_headers = ["№", "Задача", "Часы", "Сроки", "Описание работы"]
        t = make_table(doc, task_headers, WIDTHS_TASKS)
        for i, task in enumerate(member["tasks"], 1):
            add_row(t, [
                i, task["name"], task["hours_estimate"],
                f"{task['start_date']} – {task['end_date']}",
                task["work_description"],
            ], WIDTHS_TASKS)

    # 3. Встречи
    add_heading(doc, "3. Встречи с ментором в ходе итерации", level=1)
    if not data["meetings"]:
        add_text(doc, "Встреч не было.")
    else:
        t = make_table(doc, ["Дата", "Тема", "Итоги"], WIDTHS_MEETINGS)
        for m in data["meetings"]:
            add_row(t, [m["meeting_date"], m["title"], m["summary"]], WIDTHS_MEETINGS)

    # 4. Оценка ментора
    add_heading(doc, "4. Оценка ментора за итерацию", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Балл: ")
    r.bold = True
    r.font.size = Pt(12)
    r = p.add_run(f"{member['sprint_score']['score']} / 10")
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Комментарий ментора: ")
    r.bold = True
    r.font.size = Pt(12)
    r = p.add_run(member["sprint_score"]["comment"])
    r.font.size = Pt(12)

    # Подписи
    sp = doc.add_paragraph()
    _apply_para_format(sp)
    p = doc.add_paragraph()
    _apply_para_format(p)
    r = p.add_run("Подписи:")
    r.bold = True
    r.font.size = Pt(12)

    sig_w = [5.5, 5.75, 5.75]
    tbl = doc.add_table(rows=0, cols=3)
    tbl.autofit = False
    for label, name, sign in [
        ("Участник", member["full_name"], "____________________"),
        ("Ментор проекта", data["project"]["mentor"], "____________________"),
        ("Дата", data["report"]["generated_at"], ""),
    ]:
        cells = tbl.add_row().cells
        cells[0].text = label
        cells[1].text = name
        cells[2].text = sign
        for j, c in enumerate(cells):
            set_cell_width(c, sig_w[j])
            for para in c.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_para_format(para)
                for r in para.runs:
                    r.font.size = Pt(12)
                    if j == 0:
                        r.bold = True

    return doc


# ════════════════════════════════════════════════════════════
#                         MARKDOWN
# ════════════════════════════════════════════════════════════

def _md_escape(s):
    if s is None:
        return "—"
    return str(s).replace("|", "\\|")


def _md_field(label, value):
    return f"- {label}: {value if value is not None else '—'}"


def build_markdown(data):
    out = []
    push = out.append
    sprint = data["sprint"]
    project = data["project"]
    team = data["team"]
    report = data.get("report", {})

    push(f"# Отчёт по итерации №{sprint['number']} — {project['title']}\n")
    push(f"> Команда: {team['name']} · Спринт: "
         f"{sprint['start_date']} — {sprint['end_date']} · "
         f"Сформирован: {report.get('generated_at', '—')}\n")

    push("## 1. Проект\n")
    push(_md_field("ID проекта", project.get("id")))
    push(_md_field("Название", project["title"]))
    push(_md_field("Инициатор", project.get("company")))
    push(_md_field("Ментор", project.get("mentor")))
    push(_md_field("Цель", project.get("goal")))
    push(_md_field("Ожидаемый результат", project.get("expected_result")))
    push(_md_field("Технологии", ", ".join(project.get("technologies", [])) or "—"))
    push(_md_field("Требования к компетенциям", project.get("competencies")))
    push(_md_field("Критерии приёмки", project.get("acceptance_criteria")))
    push(_md_field("Образовательный результат", project.get("edu_result")))
    push(_md_field("Допустимые курсы",
                   ", ".join(str(c) for c in project.get("courses", [])) or "—"))
    push(_md_field("Минимальный средний балл", project.get("min_gpa")))
    push(_md_field("Размер команды",
                   f"{project.get('team_size_min')}–{project.get('team_size_max')}"))
    push(_md_field("Кол-во команд", project.get("num_teams")))
    push(_md_field("Срок (семестров)", project.get("duration_semesters")))
    push(_md_field("Проект-предшественник", project.get("predecessor")))
    push(_md_field("Ресурсы", project.get("resources")))
    push("")
    push("### Полное описание")
    push(project.get("full_description") or "—")
    push("")

    push("## 2. Итерация (спринт)\n")
    push(_md_field("ID спринта", sprint.get("id")))
    push(_md_field("Номер", sprint["number"]))
    push(_md_field("Дата начала", sprint["start_date"]))
    push(_md_field("Дата окончания", sprint["end_date"]))
    push(_md_field("Длительность (дней)", sprint.get("duration_days")))
    push(_md_field("Статус", sprint.get("status")))
    push("")

    push("## 3. Команда\n")
    push(_md_field("ID команды", team.get("id")))
    push(_md_field("Название", team["name"]))
    push(_md_field("Тимлид", team["leader"]))
    push(_md_field("Дата создания", team.get("created_at")))
    push("")
    push("### Состав")
    push("| # | ID | ФИО | Роль в команде | Присоединился | Группа | Курс | Ср. балл |")
    push("|---|----|-----|-----|-----|-----|-----|-----|")
    for i, m in enumerate(team["members"], 1):
        push(f"| {i} | {m.get('user_id', '—')} | {_md_escape(m['full_name'])} | "
             f"{_md_escape(m['role_in_team'])} | {m.get('joined_at', '—')} | "
             f"{m.get('group', '—')} | {m.get('course', '—')} | {m.get('gpa', '—')} |")
    push("")

    tr = data["team_report"]
    push("## 4. Командный отчёт\n")
    push(_md_field("ID отчёта", tr.get("id")))
    push(_md_field("Статус", tr.get("status")))
    push(_md_field("Отправлен на проверку", tr.get("submitted_at")))
    push(_md_field("Проверен ментором", tr.get("reviewed_at")))
    push("")
    push("### 4.1. Что сделано")
    push(tr.get("summary") or "—")
    push("")
    push("### 4.2. Проблемы и риски")
    push(tr.get("problems") or "—")
    push("")
    push("### 4.3. План на следующую итерацию")
    push(tr.get("next_plan") or "—")
    push("")
    push("### 4.4. Комментарий ментора")
    push(tr.get("mentor_comment") or "—")
    push("")

    push("## 5. Задачи по участникам\n")
    for member in data["members_with_tasks"]:
        push(f"### {member['full_name']} — {member['role_in_team']}\n")
        push(_md_field("Оценка за спринт", f"{member['sprint_score']['score']} / 10"))
        push(_md_field("Комментарий ментора к оценке", member["sprint_score"].get("comment")))
        push(_md_field("Оценку поставил", member["sprint_score"].get("scored_by")))
        push("")

        if not member["tasks"]:
            push("_Задач нет._\n")
            continue

        for i, task in enumerate(member["tasks"], 1):
            push(f"#### Задача {i}. {task['name']}")
            push(_md_field("ID задачи", task.get("id")))
            push(_md_field("Статус", task["status"]))
            push(_md_field("Просрочена когда-либо", task.get("was_overdue", False)))
            push(_md_field("Создал", task.get("created_by")))
            push(_md_field("Плановые часы", task["hours_estimate"]))
            push(_md_field("Начало", task["start_date"]))
            push(_md_field("Окончание", task["end_date"]))
            push(_md_field("Ссылка на MR", task.get("mr_link")))
            push(_md_field("Создана", task.get("created_at")))
            push(_md_field("Последнее изменение статуса", task.get("status_changed_at")))
            push("")
            push("Описание задачи:")
            push(task.get("description") or "—")
            push("")
            push("Описание выполненной работы:")
            push(task.get("work_description") or "—")
            push("")

            comments = task.get("mentor_comments", [])
            if comments:
                push("Комментарии ментора:")
                for c in comments:
                    push(f"- {c['action']}: {c['text']}")
                push("")

            history = task.get("history", [])
            if history:
                push("История событий (день спринта → событие):")
                for h in history:
                    push(f"- День {h['day']}: {h['event']}")
                push("")

    push("## 6. Встречи с ментором\n")
    if not data["meetings"]:
        push("_Встреч не было._\n")
    else:
        push("| # | ID | Дата | Время | Длит. | Тема | Статус | Кто создал | Подтв. ментором | Ссылка | Повестка | Итоги |")
        push("|---|----|------|-------|-------|------|--------|-----|------|--------|-----|-----|")
        for i, m in enumerate(data["meetings"], 1):
            push(
                f"| {i} | {m.get('id', '—')} | {m['meeting_date']} | {m['start_time']} | "
                f"{m['duration_minutes']} мин | {_md_escape(m['title'])} | "
                f"{_md_escape(m.get('status', '—'))} | {_md_escape(m.get('created_by', '—'))} | "
                f"{m.get('mentor_confirmed', '—')} | {_md_escape(m.get('conference_link', '—'))} | "
                f"{_md_escape(m.get('description', '—'))} | {_md_escape(m.get('summary', '—'))} |"
            )
        push("")

    push("## 7. Сводные индивидуальные оценки за спринт\n")
    push("| ФИО | Балл / 10 | Комментарий ментора | Кто поставил |")
    push("|-----|-----------|---------------------|-----|")
    for m in data["members_with_tasks"]:
        s = m["sprint_score"]
        push(f"| {_md_escape(m['full_name'])} | {s['score']} | "
             f"{_md_escape(s.get('comment'))} | {_md_escape(s.get('scored_by', '—'))} |")
    push("")

    push("## 8. Статистика спринта\n")
    all_tasks = [t for m in data["members_with_tasks"] for t in m["tasks"]]
    done = sum(1 for t in all_tasks if t["status"] == "Готово")
    on_review = sum(1 for t in all_tasks if t["status"] == "На ревью")
    returned = sum(1 for t in all_tasks if t["status"] == "Возвращена")
    overdue = sum(1 for t in all_tasks if t.get("was_overdue") is True)
    total_hours = sum(t["hours_estimate"] for t in all_tasks
                      if isinstance(t["hours_estimate"], (int, float)))
    push(_md_field("Всего задач", len(all_tasks)))
    push(_md_field("Готово", done))
    push(_md_field("На ревью", on_review))
    push(_md_field("Возвращено", returned))
    push(_md_field("Было просрочено", overdue))
    push(_md_field("Суммарные плановые часы", total_hours if total_hours else "—"))
    push(_md_field("Количество встреч", len(data["meetings"])))
    scores = [m["sprint_score"]["score"] for m in data["members_with_tasks"]
              if isinstance(m["sprint_score"]["score"], (int, float))]
    push(_md_field("Средний балл по команде",
                   f"{sum(scores)/len(scores):.2f}" if scores else "—"))
    push("")

    return "\n".join(out)


def build_student_markdown(data, student_index=0):
    member = data["members_with_tasks"][student_index]
    team_member = next(
        (m for m in data["team"]["members"] if m["full_name"] == member["full_name"]),
        {},
    )
    sprint = data["sprint"]
    project = data["project"]
    report = data.get("report", {})

    out = []
    push = out.append
    push(f"# Индивидуальный отчёт по итерации №{sprint['number']} — {member['full_name']}\n")
    push(f"> Проект: {project['title']} · Команда: {data['team']['name']} · "
         f"Спринт: {sprint['start_date']} — {sprint['end_date']} · "
         f"Сформирован: {report.get('generated_at', '—')}\n")

    push("## 1. Участник\n")
    push(_md_field("ID пользователя", team_member.get("user_id")))
    push(_md_field("ФИО", member["full_name"]))
    push(_md_field("Функциональная роль", member["role_in_team"]))
    push(_md_field("Группа", team_member.get("group")))
    push(_md_field("Курс", team_member.get("course")))
    push(_md_field("Средний балл", team_member.get("gpa")))
    push(_md_field("Присоединился к команде", team_member.get("joined_at")))
    push("")

    push("## 2. Проект\n")
    push(_md_field("ID проекта", project.get("id")))
    push(_md_field("Название", project["title"]))
    push(_md_field("Инициатор", project.get("company")))
    push(_md_field("Ментор", project.get("mentor")))
    push(_md_field("Команда", data["team"]["name"]))
    push(_md_field("Тимлид", data["team"]["leader"]))
    push("")

    push("## 3. Итерация (спринт)\n")
    push(_md_field("ID спринта", sprint.get("id")))
    push(_md_field("Номер", sprint["number"]))
    push(_md_field("Дата начала", sprint["start_date"]))
    push(_md_field("Дата окончания", sprint["end_date"]))
    push(_md_field("Длительность (дней)", sprint.get("duration_days")))
    push(_md_field("Статус", sprint.get("status")))
    push("")

    push("## 4. Задачи участника\n")
    if not member["tasks"]:
        push("_Задач нет._\n")
    else:
        for i, task in enumerate(member["tasks"], 1):
            push(f"### Задача {i}. {task['name']}")
            push(_md_field("ID задачи", task.get("id")))
            push(_md_field("Статус", task["status"]))
            push(_md_field("Просрочена когда-либо", task.get("was_overdue", False)))
            push(_md_field("Создал", task.get("created_by")))
            push(_md_field("Плановые часы", task["hours_estimate"]))
            push(_md_field("Начало", task["start_date"]))
            push(_md_field("Окончание", task["end_date"]))
            push(_md_field("Ссылка на MR", task.get("mr_link")))
            push(_md_field("Создана", task.get("created_at")))
            push(_md_field("Последнее изменение статуса", task.get("status_changed_at")))
            push("")
            push("Описание задачи:")
            push(task.get("description") or "—")
            push("")
            push("Описание выполненной работы:")
            push(task.get("work_description") or "—")
            push("")

            comments = task.get("mentor_comments", [])
            if comments:
                push("Комментарии ментора:")
                for c in comments:
                    push(f"- {c['action']}: {c['text']}")
                push("")

            history = task.get("history", [])
            if history:
                push("История событий (день спринта → событие):")
                for h in history:
                    push(f"- День {h['day']}: {h['event']}")
                push("")

    push("## 5. Встречи с ментором в ходе итерации\n")
    if not data["meetings"]:
        push("_Встреч не было._\n")
    else:
        push("| # | ID | Дата | Время | Длит. | Тема | Статус | Кто создал | Ссылка | Повестка | Итоги |")
        push("|---|----|------|-------|-------|------|--------|-----|--------|-----|-----|")
        for i, m in enumerate(data["meetings"], 1):
            push(
                f"| {i} | {m.get('id', '—')} | {m['meeting_date']} | {m['start_time']} | "
                f"{m['duration_minutes']} мин | {_md_escape(m['title'])} | "
                f"{_md_escape(m.get('status', '—'))} | {_md_escape(m.get('created_by', '—'))} | "
                f"{_md_escape(m.get('conference_link', '—'))} | "
                f"{_md_escape(m.get('description', '—'))} | {_md_escape(m.get('summary', '—'))} |"
            )
        push("")

    push("## 6. Оценка ментора за итерацию\n")
    s = member["sprint_score"]
    push(_md_field("Балл / 10", s["score"]))
    push(_md_field("Комментарий ментора", s.get("comment")))
    push(_md_field("Оценку поставил", s.get("scored_by")))
    push("")

    push("## 7. Личная статистика за спринт\n")
    tasks = member["tasks"]
    done = sum(1 for t in tasks if t["status"] == "Готово")
    on_review = sum(1 for t in tasks if t["status"] == "На ревью")
    returned = sum(1 for t in tasks if t["status"] == "Возвращена")
    overdue = sum(1 for t in tasks if t.get("was_overdue") is True)
    total_hours = sum(t["hours_estimate"] for t in tasks
                      if isinstance(t["hours_estimate"], (int, float)))
    push(_md_field("Всего задач", len(tasks)))
    push(_md_field("Готово", done))
    push(_md_field("На ревью", on_review))
    push(_md_field("Возвращено", returned))
    push(_md_field("Было просрочено", overdue))
    push(_md_field("Суммарные плановые часы", total_hours if total_hours else "—"))
    push("")

    return "\n".join(out)


# ════════════════════════════════════════════════════════════
#                    Данные для генерации
# ════════════════════════════════════════════════════════════

TEMPLATE_DATA = {
    "project": {
        "title": "{{project.title}}",
        "company": "{{project.company}}",
        "mentor": "{{project.mentor.full_name}}",
    },
    "team": {
        "name": "{{team.name}}",
        "leader": "{{team.leader.full_name}}",
        "members": [{"full_name": "{{member.full_name}}",
                     "role_in_team": "{{member.role_in_team}}"}],
    },
    "sprint": {
        "number": "{{sprint.number}}",
        "start_date": "{{sprint.start_date}}",
        "end_date": "{{sprint.end_date}}",
    },
    "report": {"generated_at": "{{report.generated_at}}"},
    "team_report": {
        "summary": "{{team_report.summary}}",
        "problems": "{{team_report.problems}}",
        "next_plan": "{{team_report.next_plan}}",
        "mentor_comment": "{{team_report.mentor_comment}}",
    },
    "members_with_tasks": [
        {
            "full_name": "{{member.full_name}}",
            "role_in_team": "{{member.role_in_team}}",
            "tasks": [{
                "name": "{{task.name}}",
                "status": "{{task.status}}",
                "hours_estimate": "{{task.hours_estimate}}",
                "start_date": "{{task.start_date}}",
                "end_date": "{{task.end_date}}",
                "mr_link": "{{task.mr_link}}",
                "work_description": "{{task.work_description}}",
            }],
            "sprint_score": {
                "score": "{{member.sprint_score.score}}",
                "comment": "{{member.sprint_score.comment}}",
            },
        }
    ],
    "meetings": [{
        "meeting_date": "{{meeting.meeting_date}}",
        "start_time": "{{meeting.start_time}}",
        "title": "{{meeting.title}}",
        "duration_minutes": "{{meeting.duration_minutes}}",
        "summary": "{{meeting.summary}}",
    }],
}

EXAMPLE_DATA = {
    "project": {
        "id": 1,
        "title": "Система управления проектным практикумом ВШПИ МФТИ",
        "company": "МФТИ, Высшая школа программной инженерии",
        "mentor": "Тимохин Валентин Николаевич",
        "goal": "Автоматизировать цикл управления дисциплиной «Проектный практикум» — "
                "от подачи заявок и распределения студентов до отчётности и защиты.",
        "expected_result": "Веб-система с ролями студент/тимлид/ментор/координатор, "
                           "развёрнутая на VDI МФТИ, интегрированная с Яндекс Диском.",
        "technologies": ["Go 1.24", "PostgreSQL 16", "React", "Docker", "goqu", "pgx"],
        "competencies": "Go/PostgreSQL, REST API, React, Git, Docker",
        "acceptance_criteria": "Все 5 ролей работают, документация API (OpenAPI), "
                               "покрытие тестами ≥ 70%, деплой на VDI.",
        "edu_result": "Навыки Clean Architecture, командной работы, DevOps-практик.",
        "courses": [2, 3, 4],
        "min_gpa": 6.0,
        "team_size_min": 3,
        "team_size_max": 5,
        "num_teams": 2,
        "duration_semesters": 2,
        "predecessor": None,
        "resources": "VDI-контейнер МФТИ, доступ к Яндекс Диску, консультации ментора.",
        "full_description": (
            "Веб-система заменяет ручное управление дисциплиной (Яндекс-формы, Excel). "
            "Реализуется поэтапно: сначала модули заявок и распределения, затем "
            "управление спринтами и задачами, отчётность и итоговое оценивание."
        ),
    },
    "team": {
        "id": 1,
        "name": "Команда 1",
        "leader": "Стародубов Александр Юрьевич",
        "created_at": "15.02.2026",
        "members": [
            {"user_id": 5, "full_name": "Стародубов Александр Юрьевич",
             "role_in_team": "Backend-разработчик, тимлид",
             "joined_at": "15.02.2026", "group": "Б05-321", "course": 2, "gpa": 7.8},
            {"user_id": 7, "full_name": "Кузнецов Михаил Сергеевич",
             "role_in_team": "Frontend-разработчик",
             "joined_at": "15.02.2026", "group": "Б05-322", "course": 2, "gpa": 7.1},
            {"user_id": 9, "full_name": "Лебедева Наталья Андреевна",
             "role_in_team": "Системный аналитик",
             "joined_at": "17.02.2026", "group": "Б05-321", "course": 2, "gpa": 8.4},
            {"user_id": 11, "full_name": "Волков Дмитрий Павлович",
             "role_in_team": "Инженер по тестированию",
             "joined_at": "18.02.2026", "group": "Б05-323", "course": 2, "gpa": 6.9},
        ],
    },
    "sprint": {
        "id": 2, "number": 2,
        "start_date": "17.03.2026", "end_date": "13.04.2026",
        "duration_days": 28, "status": "Завершён",
    },
    "report": {"generated_at": "14.04.2026"},
    "team_report": {
        "id": 12, "status": "Проверен",
        "submitted_at": "14.04.2026 10:22", "reviewed_at": "14.04.2026 18:45",
        "summary": (
            "Завершён этап проектирования базы данных: согласованы 7 новых таблиц "
            "(teams, team_members, sprints, tasks, team_reports, meetings, user_profiles). "
            "Реализован и развёрнут на VDI модуль авторизации (JWT + middleware ролей). "
            "Frontend: готовы страницы студента (распределён/нераспределён) и тимлида. "
            "Покрытие unit-тестами backend — 74%."
        ),
        "problems": (
            "1. Задержка с выделением VDI-контейнера МФТИ (3 дня простоя). "
            "2. Неопределённость в алгоритме чернового распределения. "
            "3. Библиотека yadisk имеет ограничения по квоте."
        ),
        "next_plan": (
            "1. Реализовать CRUD задач с аппрувом ментора. "
            "2. Сверстать диаграмму Ганта на React. "
            "3. Согласовать алгоритм распределения и реализовать сервис. "
            "4. Добавить экспорт отчётов в DOCX. "
            "5. Интеграционные тесты для модуля команд."
        ),
        "mentor_comment": (
            "Хороший темп работы. Отмечаю качество схемы БД — предусмотрены soft-delete "
            "и триггеры для уведомлений. К следующему спринту — OpenAPI-спецификация."
        ),
    },
    "members_with_tasks": [
        {
            "full_name": "Стародубов Александр Юрьевич",
            "role_in_team": "Backend-разработчик, тимлид",
            "tasks": [
                {"id": 101, "name": "Схема БД: 7 новых таблиц + миграции",
                 "description": "Спроектировать и применить миграции для teams, team_members, sprints, tasks, team_reports, meetings, user_profiles.",
                 "status": "Готово", "was_overdue": False, "created_by": "Стародубов А.Ю.",
                 "hours_estimate": 16, "start_date": "17.03", "end_date": "23.03",
                 "mr_link": "!42",
                 "work_description": "Применены SQL-миграции 7 таблиц, триггеры update_updated_at и update_status_changed_at, индексы под запросы уведомлений.",
                 "mentor_comments": [
                     {"action": "Аппрув", "text": "Задача понятная, объём адекватен."},
                     {"action": "Принятие", "text": "Отличная работа, триггеры оценил."}],
                 "history": [{"day": 6, "event": "review"}, {"day": 7, "event": "accepted"}],
                 "created_at": "17.03.2026 09:12", "status_changed_at": "24.03.2026 14:05"},
                {"id": 104, "name": "Модуль авторизации: JWT + middleware ролей",
                 "description": "Реализовать генерацию/валидацию JWT и middleware для 5 ролей.",
                 "status": "Готово", "was_overdue": False, "created_by": "Стародубов А.Ю.",
                 "hours_estimate": 12, "start_date": "24.03", "end_date": "30.03",
                 "mr_link": "!45",
                 "work_description": "JWT + middleware на 5 ролей, unit-тесты позитив/негатив.",
                 "mentor_comments": [
                     {"action": "Аппрув", "text": "Можно приступать."},
                     {"action": "Принятие", "text": "Принято."}],
                 "history": [{"day": 13, "event": "review"}, {"day": 14, "event": "accepted"}],
                 "created_at": "24.03.2026 10:00", "status_changed_at": "31.03.2026 16:20"},
                {"id": 110, "name": "CRUD команд и участников (teams API)",
                 "description": "POST/GET/PUT/DELETE /api/teams + /members с валидацией прав.",
                 "status": "На ревью", "was_overdue": False, "created_by": "Стародубов А.Ю.",
                 "hours_estimate": 10, "start_date": "31.03", "end_date": "10.04",
                 "mr_link": "!51",
                 "work_description": "Все 4 ручки + интеграционные тесты на pgx.",
                 "mentor_comments": [{"action": "Аппрув", "text": "ОК."}],
                 "history": [{"day": 25, "event": "review"}],
                 "created_at": "31.03.2026 11:30", "status_changed_at": "11.04.2026 09:10"},
            ],
            "sprint_score": {
                "score": 9,
                "comment": "Отличная архитектура. Снят 1 балл за отсутствие OpenAPI.",
                "scored_by": "Тимохин В.Н."},
        },
        {
            "full_name": "Кузнецов Михаил Сергеевич",
            "role_in_team": "Frontend-разработчик",
            "tasks": [
                {"id": 102, "name": "Страница нераспределённого студента (каталог, 5 приоритетов)",
                 "description": "Сверстать каталог проектов + слоты приоритетов с DnD.",
                 "status": "Готово", "was_overdue": False, "created_by": "Кузнецов М.С.",
                 "hours_estimate": 14, "start_date": "17.03", "end_date": "26.03",
                 "mr_link": "!43",
                 "work_description": "Drag-and-drop, FLIP-анимация, inline-подробности, /api/projects.",
                 "mentor_comments": [
                     {"action": "Аппрув", "text": "Ок."},
                     {"action": "Принятие", "text": "Аккуратная реализация."}],
                 "history": [{"day": 9, "event": "review"}, {"day": 10, "event": "accepted"}],
                 "created_at": "17.03.2026 09:30", "status_changed_at": "27.03.2026 15:40"},
                {"id": 107, "name": "Страница распределённого студента + диаграмма Ганта (базовая)",
                 "description": "Реализовать Гант со sticky-колонками и попапом редактирования.",
                 "status": "Готово", "was_overdue": False, "created_by": "Кузнецов М.С.",
                 "hours_estimate": 18, "start_date": "27.03", "end_date": "10.04",
                 "mr_link": "!48",
                 "work_description": "Sticky-колонки, попап редактирования, цветные линии истории.",
                 "mentor_comments": [
                     {"action": "Аппрув", "text": "ОК."},
                     {"action": "Принятие", "text": "Принято."}],
                 "history": [{"day": 22, "event": "review"}, {"day": 23, "event": "accepted"}],
                 "created_at": "27.03.2026 10:00", "status_changed_at": "11.04.2026 12:00"},
            ],
            "sprint_score": {
                "score": 8,
                "comment": "Качественная вёрстка. К доработке: адаптивность < 1280px.",
                "scored_by": "Тимохин В.Н."},
        },
        {
            "full_name": "Лебедева Наталья Андреевна",
            "role_in_team": "Системный аналитик",
            "tasks": [
                {"id": 103, "name": "Актуализация architecture.md до версии 1.1",
                 "description": "Описать ДКА задач/заявок, схему проектов-продолжений.",
                 "status": "Готово", "was_overdue": False, "created_by": "Лебедева Н.А.",
                 "hours_estimate": 8, "start_date": "17.03", "end_date": "22.03",
                 "mr_link": "!44",
                 "work_description": "11 статусов ДКА, правила распределения, согласовано с ментором.",
                 "mentor_comments": [
                     {"action": "Аппрув", "text": "Делаем."},
                     {"action": "Принятие", "text": "Очень чисто."}],
                 "history": [{"day": 5, "event": "review"}, {"day": 6, "event": "accepted"}],
                 "created_at": "17.03.2026 09:45", "status_changed_at": "23.03.2026 17:00"},
                {"id": 111, "name": "Проект регламента docx-отчётов + шаблон",
                 "description": "Проанализировать Положение и разработать шаблон отчёта.",
                 "status": "На ревью", "was_overdue": False, "created_by": "Лебедева Н.А.",
                 "hours_estimate": 6, "start_date": "07.04", "end_date": "13.04",
                 "mr_link": "!52",
                 "work_description": "5 разделов, два формата выгрузки (DOCX + MD), пример заполнения.",
                 "mentor_comments": [{"action": "Аппрув", "text": "Хорошая идея разделения форматов."}],
                 "history": [{"day": 28, "event": "review"}],
                 "created_at": "07.04.2026 14:00", "status_changed_at": "13.04.2026 19:00"},
            ],
            "sprint_score": {
                "score": 9,
                "comment": "Архитектурная документация на высоком уровне.",
                "scored_by": "Тимохин В.Н."},
        },
        {
            "full_name": "Волков Дмитрий Павлович",
            "role_in_team": "Инженер по тестированию",
            "tasks": [
                {"id": 106, "name": "Интеграционные тесты для auth-модуля",
                 "description": "testcontainers + PostgreSQL, все ролевые сценарии.",
                 "status": "Готово", "was_overdue": False, "created_by": "Волков Д.П.",
                 "hours_estimate": 10, "start_date": "25.03", "end_date": "02.04",
                 "mr_link": "!47",
                 "work_description": "28 тестов, coverage 91% по auth-модулю.",
                 "mentor_comments": [
                     {"action": "Аппрув", "text": "ОК."},
                     {"action": "Принятие", "text": "Хорошее покрытие."}],
                 "history": [{"day": 15, "event": "review"}, {"day": 16, "event": "accepted"}],
                 "created_at": "25.03.2026 11:00", "status_changed_at": "03.04.2026 10:00"},
                {"id": 112, "name": "Настройка CI: линтер + тесты на PR",
                 "description": "GitLab CI с golangci-lint и go test.",
                 "status": "Возвращена", "was_overdue": True, "created_by": "Волков Д.П.",
                 "hours_estimate": 6, "start_date": "03.04", "end_date": "10.04",
                 "mr_link": "!50",
                 "work_description": "Базовая конфигурация готова, не хватает стадии Docker build.",
                 "mentor_comments": [
                     {"action": "Аппрув", "text": "Делаем."},
                     {"action": "Возврат", "text": "Добавь стадию Docker build и publish."}],
                 "history": [{"day": 24, "event": "review"}, {"day": 26, "event": "returned"}],
                 "created_at": "03.04.2026 10:00", "status_changed_at": "12.04.2026 15:00"},
            ],
            "sprint_score": {
                "score": 7,
                "comment": "Хорошие интеграционные тесты. Доработать CI с Docker build.",
                "scored_by": "Тимохин В.Н."},
        },
    ],
    "meetings": [
        {"id": 21, "meeting_date": "19.03.2026", "start_time": "17:00",
         "title": "Планирование спринта 2",
         "description": "Повестка: распределение задач, приоритет на БД и auth.",
         "duration_minutes": 60,
         "conference_link": "https://telemost.yandex.ru/j/2026031917",
         "summary": "Задачи спринта согласованы и распределены.",
         "status": "Состоялась", "created_by": "Стародубов А.Ю. (тимлид)",
         "mentor_confirmed": True},
        {"id": 25, "meeting_date": "02.04.2026", "start_time": "17:00",
         "title": "Мид-спринт синк",
         "description": "Повестка: блокеры, обсуждение алгоритма распределения.",
         "duration_minutes": 45,
         "conference_link": "https://telemost.yandex.ru/j/2026040217",
         "summary": "Алгоритм распределения — решено обсудить отдельно с Координатором ПП.",
         "status": "Состоялась", "created_by": "Тимохин В.Н. (ментор)",
         "mentor_confirmed": True},
        {"id": 29, "meeting_date": "13.04.2026", "start_time": "17:00",
         "title": "Обзор спринта 2 + планирование спринта 3",
         "description": "Демо готовых модулей + план на спринт 3.",
         "duration_minutes": 90,
         "conference_link": "https://telemost.yandex.ru/j/2026041317",
         "summary": "Приняты: БД, auth, Гант. В спринт 3 — экспорт DOCX и алгоритм распределения.",
         "status": "Состоялась", "created_by": "Тимохин В.Н. (ментор)",
         "mentor_confirmed": True},
    ],
}


if __name__ == "__main__":
    out = "/home/alex/MIPT_Project_Management_System"
    build_docx(TEMPLATE_DATA).save(f"{out}/sprint_report_template.docx")
    build_docx(EXAMPLE_DATA).save(f"{out}/sprint_report_example.docx")
    with open(f"{out}/sprint_report_template.md", "w", encoding="utf-8") as f:
        f.write(build_markdown(TEMPLATE_DATA))
    with open(f"{out}/sprint_report_example.md", "w", encoding="utf-8") as f:
        f.write(build_markdown(EXAMPLE_DATA))
    build_student_docx(TEMPLATE_DATA).save(f"{out}/sprint_report_student_template.docx")
    build_student_docx(EXAMPLE_DATA).save(f"{out}/sprint_report_student_example.docx")
    with open(f"{out}/sprint_report_student_template.md", "w", encoding="utf-8") as f:
        f.write(build_student_markdown(TEMPLATE_DATA))
    with open(f"{out}/sprint_report_student_example.md", "w", encoding="utf-8") as f:
        f.write(build_student_markdown(EXAMPLE_DATA))
    print("saved: team/student × template/example × (docx + md) = 8 файлов")
